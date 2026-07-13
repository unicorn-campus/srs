# -*- coding: utf-8 -*-
"""MD -> DOCX 변환 (LG유플러스 개발요구사항정의서 양식 재현).
A4 / 여백 상3.0·좌우하2.54cm / 맑은 고딕 / 제목·소제목 볼드 / 글머리표 / 표 4종 머리행 음영 / ⚠ 볼드.
"""
import sys, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:/Users/hiond/class/srs/output/솔루션요구사항정의서.md"
DST = r"C:/Users/hiond/class/srs/output/솔루션요구사항정의서.docx"
FONT = "Malgun Gothic"

def set_run(run, size=None, bold=None, color=None):
    run.font.name = FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'):
        rFonts.set(qn(a), FONT)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if color is not None: run.font.color.rgb = color

def add_runs(p, text, size, bold=False, color=None):
    # inline **bold** 처리
    for i, seg in enumerate(text.split('**')):
        if seg == '': continue
        r = p.add_run(seg)
        set_run(r, size=size, bold=(bold or (i % 2 == 1)), color=color)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
sec.top_margin = Cm(3.0); sec.bottom_margin = Cm(2.54)
sec.left_margin = Cm(2.54); sec.right_margin = Cm(2.54)
normal = doc.styles['Normal']
normal.font.name = FONT; normal.font.size = Pt(10.5)
normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), FONT)

lines = open(SRC, encoding='utf-8').read().split('\n')
n = len(lines); i = 0
seen_chapter = False

def col_widths(header):
    ncol = len(header)
    if ncol == 2: return [4.0, 11.9]
    if ncol == 3: return [3.5, 6.2, 6.2]
    if ncol == 4:
        if header[0].strip() == 'No': return [1.0, 3.2, 7.8, 3.9]
        return [2.6, 3.9, 5.4, 4.0]  # 상태값 표
    return [15.9/ncol]*ncol

def add_table(block):
    rows = []
    for ln in block:
        cells = [c.strip() for c in ln.strip().strip('|').split('|')]
        rows.append(cells)
    # rows[1] = 구분선 제거
    header = rows[0]; body = rows[2:]
    data = [header] + body
    w = col_widths(header)
    t = doc.add_table(rows=len(data), cols=len(header))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; t.allow_autofit = False
    for ri, row in enumerate(data):
        for ci in range(len(header)):
            cell = t.cell(ri, ci)
            cell.width = Cm(w[ci])
            p = cell.paragraphs[0]
            for r in list(p.runs): r._element.getparent().remove(r._element)
            txt = row[ci] if ci < len(row) else ''
            add_runs(p, txt, 9.5, bold=(ri == 0))
            if ri == 0: shade(cell, 'D9D9D9')
    doc.add_paragraph()

while i < n:
    line = lines[i]
    s = line.strip()
    if s == '' or s == '---':
        i += 1; continue
    # 표
    if s.startswith('|'):
        block = []
        while i < n and lines[i].strip().startswith('|'):
            block.append(lines[i]); i += 1
        if len(block) >= 2: add_table(block)
        continue
    # 소제목 ###
    if s.startswith('### '):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
        add_runs(p, s[4:], 11.5, bold=True)
        i += 1; continue
    # 장 제목 ##
    if s.startswith('## '):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
        add_runs(p, s[3:], 13, bold=True)
        seen_chapter = True; i += 1; continue
    # 글머리표
    if s.startswith('- '):
        txt = s[2:]
        j = i + 1
        while j < n and lines[j].startswith('  ') and lines[j].strip() != '':
            txt += ' ' + lines[j].strip(); j += 1
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        bold = txt.startswith('⚠')
        add_runs(p, txt, 10.5, bold=bold)
        i = j; continue
    # 일반 문단 (연속 비공백·비특수 줄 병합)
    para = [s]; j = i + 1
    while j < n:
        nx = lines[j]; nxs = nx.strip()
        if nxs == '' or nxs == '---': break
        if nxs.startswith(('#', '|', '- ')): break
        para.append(nxs); j += 1
    text = ' '.join(para)
    # 머리말 처리
    if not seen_chapter and text.startswith('**문서.'):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        add_runs(p, text.strip('*'), 16, bold=True)
    elif not seen_chapter and text.startswith('※'):
        p = doc.add_paragraph(); add_runs(p, text, 9, color=RGBColor(0x66,0x66,0x66))
    elif not seen_chapter and text.startswith('MVNO'):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8); add_runs(p, text, 10.5, bold=True)
    elif text.startswith('※'):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8)
        add_runs(p, text, 9, color=RGBColor(0x66,0x66,0x66))
    else:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
        bold = text.startswith('⚠')
        add_runs(p, text, 10.5, bold=bold)
    i = j

doc.save(DST)
print("SAVED:", DST)
d2 = Document(DST)
print("paragraphs:", len(d2.paragraphs), "| tables:", len(d2.tables))
for ti, t in enumerate(d2.tables, 1):
    print(f"  table{ti}: {len(t.rows)}x{len(t.columns)} head={[c.text for c in t.rows[0].cells]}")
